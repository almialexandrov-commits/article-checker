import os
import re
import uuid
import requests
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

app = FastAPI()

os.makedirs("uploads", exist_ok=True)
os.makedirs("processed", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/", response_class=HTMLResponse)
async def read_index():
    with open("static/index.html", "r", encoding="utf-8") as f:
        return f.read()

@app.post("/check")
async def check_article(file: UploadFile = File(...)):
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Принимаются только файлы в формате .docx")

    file_id = str(uuid.uuid4())
    input_path = f"uploads/{file_id}.docx"
    output_path = f"processed/{file_id}_checked.docx"
    
    with open(input_path, "wb") as buffer:
        buffer.write(await file.read())

    try:
        doc = Document(input_path)
        report, doc = analyze_and_modify_document(doc, file.filename)
        doc.save(output_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка обработки: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.remove(input_path)

    return {
        "filename": file.filename,
        "download_url": f"/download/{file_id}",
        "report": report
    }

@app.get("/download/{file_id}")
async def download_file(file_id: str):
    file_path = f"processed/{file_id}_checked.docx"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Файл не найден")
    return FileResponse(file_path, filename="article_checked.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def add_word_comment(paragraph, comment_text, comment_id):
    """
    Добавляет нативный комментарий Word к абзацу через lxml
    """
    try:
        # Получаем доступ к части документа с комментариями
        comments_part = paragraph.part.related_parts.get('comments')
        if comments_part is None:
            # Если части комментариев нет, создаем её
            from docx.opc.constants import RELATIONSHIP_TYPE
            comments_part = paragraph.part.package.part_related_parts.get(RELATIONSHIP_TYPE.COMMENTS)
            if comments_part is None:
                # Не удалось создать часть комментариев, пропускаем
                return comment_id + 1
        
        # Создаем XML элемент комментария
        comment_xml = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment', {
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id': str(comment_id),
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author': 'Автоматическая проверка',
        })
        
        # Добавляем текст комментария
        p_xml = etree.SubElement(comment_xml, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
        r_xml = etree.SubElement(p_xml, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        t_xml = etree.SubElement(r_xml, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        t_xml.text = comment_text
        
        # Добавляем комментарий в документ
        comments_part.element.append(comment_xml)
        
        # Добавляем ссылку на комментарий в абзац
        p_element = paragraph._element
        comment_range_start = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeStart', {
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id': str(comment_id)
        })
        p_element.insert(0, comment_range_start)
        
        comment_range_end = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeEnd', {
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id': str(comment_id)
        })
        p_element.append(comment_range_end)
        
        comment_ref = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        comment_ref_xml = etree.SubElement(comment_ref, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentReference', {
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id': str(comment_id)
        })
        p_element.append(comment_ref)
        
        return comment_id + 1
    except Exception as e:
        print(f"Error adding comment: {e}")
        return comment_id + 1

def analyze_and_modify_document(doc: Document, filename: str) -> tuple:
    errors = []
    checklist = []
    paragraphs = doc.paragraphs
    full_text = "\n".join([p.text.strip() for p in paragraphs])
    comment_id = 1
    
    # --- 1. Структура и порядок ---
    indices = {
        'udk': -1, 'author': -1, 'affiliation': -1, 'title': -1,
        'abstract': -1, 'keywords': -1, 'citation': -1, 'sources': -1
    }
    
    for i, p in enumerate(paragraphs):
        text = p.text.lower().strip()
        if re.match(r'^удк\s+\d', text): indices['udk'] = i
        if "аннотация" in text and indices['abstract'] == -1: indices['abstract'] = i
        if "ключевые слова" in text and indices['keywords'] == -1: indices['keywords'] = i
        # ИСПРАВЛЕНО: ищем и "Для цитирования", и просто "Цитирование"
        if ("для цитирования" in text or text.startswith("цитирование")) and indices['citation'] == -1: 
            indices['citation'] = i
        if "список источников" in text or "список литературы" in text: 
            if indices['sources'] == -1: indices['sources'] = i

    # УДК
    if indices['udk'] != -1:
        checklist.append({"item": "УДК указан", "status": "success", "details": "Найден"})
    else:
        errors.append({"type": "error", "category": "Структура", "msg": "Не найден индекс УДК в начале статьи."})
        checklist.append({"item": "УДК указан", "status": "error", "details": "Не найден"})

    # DOI
    doi_match = re.search(r'https?://doi\.org/10\.24412/2224-5391-\d{4}-\d{2}-\d+', full_text)
    if doi_match:
        checklist.append({"item": "DOI (формат журнала)", "status": "success", "details": "Найден"})
    else:
        errors.append({"type": "warning", "category": "Метаданные", "msg": "DOI не найден или не соответствует формату (должен начинаться с https://doi.org/10.24412/2224-5391-...)"})
        checklist.append({"item": "DOI (формат журнала)", "status": "warning", "details": "Не найден или неверный формат"})

    # EDN
    if re.search(r'EDN\s+([А-ЯA-Z]{6}|Х{6})', full_text, re.IGNORECASE):
        checklist.append({"item": "EDN (рыба)", "status": "success", "details": "Найден"})
    else:
        errors.append({"type": "warning", "category": "Метаданные", "msg": "Не найдена запись EDN (например, EDN ХХХХХХ)."})
        checklist.append({"item": "EDN (рыба)", "status": "warning", "details": "Не найдена"})

    # Порядок элементов
    order_ok = True
    if indices['abstract'] != -1 and indices['keywords'] != -1 and indices['abstract'] > indices['keywords']:
        order_ok = False
        errors.append({"type": "error", "category": "Структура", "msg": "Нарушен порядок: 'Ключевые слова' идут перед 'Аннотацией'."})
        comment_id = add_word_comment(paragraphs[indices['keywords']], "Порядок нарушен: Аннотация должна быть раньше ключевых слов", comment_id)
    
    if order_ok and indices['abstract'] != -1:
        checklist.append({"item": "Порядок элементов", "status": "success", "details": "Корректный"})
    elif indices['abstract'] == -1:
        errors.append({"type": "error", "category": "Структура", "msg": "Не найдена 'Аннотация'."})
        checklist.append({"item": "Порядок элементов", "status": "error", "details": "Аннотация не найдена"})

    # Объем аннотации
    if indices['abstract'] != -1 and indices['keywords'] != -1:
        abs_text = ""
        for i in range(indices['abstract'], indices['keywords']):
            abs_text += paragraphs[i].text
        abs_len = len(abs_text.replace(" ", "").replace("\n", ""))
        if 2000 <= abs_len <= 2200:
            checklist.append({"item": "Аннотация (2000-2200 зн.)", "status": "success", "details": f"{abs_len} знаков"})
        else:
            errors.append({"type": "warning", "category": "Структура", "msg": f"Объем аннотации: {abs_len} знаков. Требуется 2000-2200."})
            checklist.append({"item": "Аннотация (2000-2200 зн.)", "status": "warning", "details": f"{abs_len} знаков"})

    # Ключевые слова (ИСПРАВЛЕНО!)
    if indices['keywords'] != -1:
        kw_text = ""
        end_idx = indices['citation'] if indices['citation'] != -1 else len(paragraphs)
        
        # Собираем текст только до следующего раздела
        for i in range(indices['keywords'], end_idx):
            p_text = paragraphs[i].text.strip()
            # Останавливаемся, если нашли следующий раздел
            if "цитирование" in p_text.lower() or "для цитирования" in p_text.lower():
                break
            kw_text += p_text + " "
        
        # Извлекаем только текст после "ключевые слова"
        kw_match = re.search(r'ключевые\s+слова\s*[:\.]?\s*(.+)', kw_text, re.IGNORECASE)
        if kw_match:
            kw_text = kw_match.group(1)
        
        # Разбиваем по запятым и считаем
        kw_list = [kw.strip() for kw in kw_text.split(',') if kw.strip() and len(kw.strip()) > 2]
        
        if 5 <= len(kw_list) <= 10:
            checklist.append({"item": "Ключевые слова (5-10, через запятую)", "status": "success", "details": f"{len(kw_list)} слов"})
        else:
            errors.append({"type": "warning", "category": "Структура", "msg": f"Ключевых слов: {len(kw_list)}. Требуется 5-10, разделенных запятыми."})
            checklist.append({"item": "Ключевые слова (5-10, через запятую)", "status": "warning", "details": f"{len(kw_list)} слов"})
            if indices['keywords'] != -1:
                comment_id = add_word_comment(paragraphs[indices['keywords']], f"Ключевых слов: {len(kw_list)}. Требуется 5-10.", comment_id)
    else:
        checklist.append({"item": "Ключевые слова (5-10, через запятую)", "status": "error", "details": "Не найдены"})

    # Цитирование и название журнала (ИСПРАВЛЕНО!)
    if indices['citation'] != -1:
        cit_text = ""
        for i in range(indices['citation'], min(indices['citation']+3, len(paragraphs))):
            cit_text += paragraphs[i].text
        if "Вестник Екатеринбургской духовной семинарии" in cit_text or "Vestnik Ekaterinburgskoi dukhovnoi seminarii" in cit_text:
            checklist.append({"item": "Название журнала в цитировании", "status": "success", "details": "Верно"})
        else:
            errors.append({"type": "error", "category": "Цитирование", "msg": "В блоке 'Цитирование' не указано название 'Вестник Екатеринбургской духовной семинарии'."})
            comment_id = add_word_comment(paragraphs[indices['citation']], "Укажите правильное название журнала", comment_id)
            checklist.append({"item": "Название журнала в цитировании", "status": "error", "details": "Неверно или не найдено"})
    else:
        errors.append({"type": "error", "category": "Структура", "msg": "Не найден блок 'Цитирование' или 'Для цитирования'."})
        checklist.append({"item": "Цитирование", "status": "error", "details": "Не найден"})

    # Даты поступления
    if re.search(r'Поступила в редакцию \d{2}\.\d{2}\.\d{4}', full_text) and re.search(r'Принята к публикации \d{2}\.\d{2}\.\d{4}', full_text):
        checklist.append({"item": "Даты поступления/принятия (рыба)", "status": "success", "details": "Найдены"})
    else:
        errors.append({"type": "warning", "category": "Метаданные", "msg": "Не найдены строки 'Поступила в редакцию 00.00.0000' и 'Принята к публикации 00.00.0000'."})
        checklist.append({"item": "Даты поступления/принятия (рыба)", "status": "warning", "details": "Не найдены"})

    # --- 2. Типографика ---
    font_issues = []
    for i, p in enumerate(paragraphs):
        if len(p.text.strip()) > 50:
            for run in p.runs:
                if not run.text.strip() or is_church_slavonic_or_greek(run.text):
                    continue
                if run.font.name and run.font.name not in ['Times New Roman', 'Times', None]:
                    font_issues.append(f"шрифт '{run.font.name}'")
                    break
                if run.font.size and run.font.size != Pt(12):
                    font_issues.append(f"размер {run.font.size.pt} пт")
                    break
    
    if not font_issues:
        checklist.append({"item": "Шрифт Times New Roman 12 пт", "status": "success", "details": "Соответствует"})
    else:
        unique_issues = list(set(font_issues))[:2]
        errors.append({"type": "warning", "category": "Типографика", "msg": f"Несоответствие шрифта/размера (например: {', '.join(unique_issues)}). Исключение: греческий/церковнославянский."})
        checklist.append({"item": "Шрифт Times New Roman 12 пт", "status": "warning", "details": f"Найдено: {', '.join(unique_issues)}"})

    # Инициалы и неразрывные пробелы
    if re.search(r'[А-Я]\.[А-Я]\.', full_text):
        errors.append({"type": "warning", "category": "Типографика", "msg": "Между инициалами обнаружены обычные пробелы или их отсутствие. Рекомендуется использовать неразрывный пробел (Ctrl+Shift+Space)."})
        checklist.append({"item": "Неразрывные пробелы в инициалах", "status": "warning", "details": "Найдены обычные пробелы"})
    else:
        checklist.append({"item": "Неразрывные пробелы в инициалах", "status": "success", "details": "Проверено"})

    # Тире
    long_dash_in_bib = re.search(r'С\.\s*\d+\s*—\s*\d+', full_text)
    if long_dash_in_bib:
        errors.append({"type": "warning", "category": "Типографика", "msg": "В диапазоне страниц обнаружено длинное тире (—). По правилам должно быть среднее тире (–)."})
        checklist.append({"item": "Среднее тире в диапазонах", "status": "warning", "details": "Найдено длинное тире"})
    else:
        checklist.append({"item": "Среднее тире в диапазонах", "status": "success", "details": "Соответствует"})

    # --- 3. Сноски ---
    footnote_refs = doc.element.body.findall('.//' + qn('w:footnoteReference'))
    if len(footnote_refs) > 0:
        checklist.append({"item": "Наличие сносок в тексте", "status": "success", "details": f"Найдено {len(footnote_refs)}"})
        endnotes = doc.element.body.findall('.//' + qn('w:endnoteReference'))
        if len(endnotes) > 0:
            errors.append({"type": "error", "category": "Сноски", "msg": "Обнаружены концевые сноски. Требуются постраничные (Вставка -> Сноска)."})
            checklist.append({"item": "Тип сносок (постраничные)", "status": "error", "details": "Найдены концевые"})
        else:
            checklist.append({"item": "Тип сносок (постраничные)", "status": "success", "details": "Соответствует"})
    else:
        errors.append({"type": "warning", "category": "Сноски", "msg": "В тексте не найдены автоматические сноски. Убедитесь, что они расставлены через меню 'Вставка -> Сноска'."})
        checklist.append({"item": "Наличие сносок в тексте", "status": "warning", "details": "Не найдены"})

    # --- 4. Богословская специфика и Орфография ---
    errors.append({"type": "info", "category": "Богословие", "msg": "Проверьте общепринятые сокращения имен святых и терминов: https://azbyka.ru/otechnik/Spravochniki/obsheprinjatye-sokrashenija-tserkovnyh-terminov/"})
    
    # Проверка формата ссылок на Библию
    bible_bad_format = re.search(r'\([А-Я][а-я]+\.\d+:\d+', full_text)
    if bible_bad_format:
        errors.append({"type": "warning", "category": "Богословие", "msg": "Обнаружен возможный неверный формат ссылки на Библию (например, 'Мф.1:1'). Правильный формат: '(Мф 1. 18–20)'."})

    # Яндекс Спеллер
    main_text = " ".join([p.text for p in paragraphs if len(p.text.strip()) > 50])
    spelling_errors = []
    if main_text:
        try:
            url = "https://speller.yandex.net/services/spellservice.json/checkText"
            response = requests.post(url, data={"text": main_text[:10000], "lang": "ru"}, timeout=5)
            results = response.json()
            for res in results:
                word = res.get('word', '')
                if any(abbr in word for abbr in ['Мф', 'Лк', 'Ин', 'Быт', 'Исх']) or is_church_slavonic_or_greek(word):
                    continue
                context = res.get('context', {}).get('text', '')
                if is_church_slavonic_or_greek(context):
                    continue
                spelling_errors.append(f"'{word}' -> {', '.join(res.get('s', [])[:2])}")
            
            if spelling_errors:
                errors.append({"type": "warning", "category": "Орфография", "msg": f"Возможные ошибки: {', '.join(spelling_errors[:5])}... (Церковнославянские слова игнорируются, но проверьте их вручную)."})
        except Exception:
            pass

    # --- 5. Генерация отчета в конце документа ---
    doc.add_page_break()
    h = doc.add_heading('📝 ОТЧЕТ АВТОМАТИЧЕСКОЙ ПРОВЕРКИ (Вестник ЕДС)', level=1)
    
    if not errors:
        p = doc.add_paragraph("✅ Грубых нарушений правил оформления не выявлено.")
    else:
        for err in errors:
            icon = "❌" if err['type'] == 'error' else ("⚠️" if err['type'] == 'warning' else "ℹ️")
            p = doc.add_paragraph(f"{icon} [{err['category'].upper()}] {err['msg']}")

    report = {
        "is_valid": not any(e['type'] == 'error' for e in errors),
        "errors": errors,
        "checklist": checklist
    }
    return report, doc

def is_church_slavonic_or_greek(text):
    if re.search(r'[\u0370-\u03FF\u1F00-\u1FFF]', text): return True
    if re.search(r'[ѣѢѳѲѵѴѧѩѫѭѯѱ\u0480-\u04FF]', text): return True
    if re.search(r'\b\w+(овъ|евъ|інъ|їнъ|іе|їе|h)\b', text, re.IGNORECASE): return True
    return False

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
